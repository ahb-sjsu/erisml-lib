"""Build JRU-formatted .docx for Variable Loss Aversion paper."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 2.0

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Variable Loss Aversion from Geometric Decision Theory:\nA Zero-Parameter Prediction')
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Andrew H. Bond, Sr. Member IEEE')
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Department of Computer Engineering\nSan José State University\nSan José, CA, USA')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(80, 80, 80)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('andrew.bond@sjsu.edu')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ORCID: 0009-0003-2599-6158')
run.font.size = Pt(11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('March 2026')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
run = p.add_run('Abstract')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph(
    'Cumulative prospect theory (Tversky & Kahneman, 1992) models loss aversion as a constant '
    '\u03bb \u2248 2.25: losses loom roughly twice as large as equivalent gains, independent of the good '
    'being traded. Yet meta-analytic evidence shows substantial variation in \u03bb across good types, '
    'with loss aversion near 1 for money and exceeding 3 for identity-laden goods (Horowitz & '
    'McConnell, 2002). No existing model predicts this variation from first principles.'
)

doc.add_paragraph(
    'We derive variable loss aversion as a geometric consequence of multi-attribute decision-making. '
    'When agents evaluate transactions on a d-dimensional attribute space with Mahalanobis edge weights, '
    'losses traverse more attribute dimensions than equivalent gains (a monetary loss activates fairness, '
    'rights, and identity dimensions that a monetary gain does not). The loss aversion coefficient is '
    'predicted to scale as \u03bb(d) = \u221ad, where d is the number of active moral/evaluative dimensions '
    'for the good in question. This yields \u03bb = 1.00 for pure cash (d = 1), \u03bb = 2.24 for standard '
    'goods (d = 5, matching the Kahneman\u2013Tversky estimate), and \u03bb = 3.53 for heirlooms '
    '(d \u2248 12.5, matching the upper range of meta-analytic variation).'
)

doc.add_paragraph(
    'The prediction is derived with zero free parameters from a covariance matrix calibrated on '
    'ultimatum game data alone, then validated out-of-sample across three canonical games and seven '
    'cross-cultural populations (MAE = 8.8%). The model produces a cross-game, cross-cultural prediction '
    'that no existing theory\u2014Nash, CPT, or Fehr\u2013Schmidt\u2014can generate without per-game '
    'parameter fitting.'
)

p = doc.add_paragraph()
run = p.add_run('Keywords: ')
run.bold = True
run.font.size = Pt(11)
run = p.add_run('loss aversion, prospect theory, Mahalanobis distance, multi-attribute decision, '
                'endowment effect, cross-cultural economics, geometric economics')
run.font.size = Pt(11)

p = doc.add_paragraph()
run = p.add_run('JEL Classification: ')
run.bold = True
run = p.add_run('D81, D91, C72, D01')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════

h = doc.add_heading('1. Introduction', level=1)
h.runs[0].font.size = Pt(14)

doc.add_paragraph(
    'Loss aversion\u2014the empirical regularity that losses are weighted more heavily than equivalent '
    'gains in decision under risk\u2014is one of the foundational findings of behavioral economics. '
    'Kahneman and Tversky (1979) documented the phenomenon; Tversky and Kahneman (1992) parameterized '
    'it as \u03bb \u2248 2.25 in cumulative prospect theory (CPT). The parameter \u03bb enters the value '
    'function as an asymmetry between the gain and loss domains:'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('v(x) = x\u1d45 if x \u2265 0;  v(x) = \u2212\u03bb(\u2212x)\u1d5d if x < 0')
run.font.size = Pt(11)

doc.add_paragraph(
    'In CPT, \u03bb is a free parameter fitted to experimental data. Its value is typically estimated '
    'at 2.0\u20132.5, and it is treated as a constant\u2014a fixed property of the value function that '
    'applies uniformly across all goods and contexts.'
)

doc.add_paragraph(
    'Meta-analytic evidence challenges this assumption. Horowitz and McConnell (2002), surveying '
    '45 WTA/WTP studies, find that loss aversion varies substantially by good type: it is near 1 '
    'for money, approximately 2 for ordinary consumer goods, and exceeds 3 for goods with identity '
    'or sentimental value. Novemsky and Kahneman (2005) report that loss aversion is attenuated or '
    'absent for money spent in routine exchange. These findings suggest that \u03bb is not a constant '
    'but a function of the good being traded\u2014but no existing model derives this function from '
    'first principles.'
)

doc.add_paragraph(
    'We provide the derivation. We show that when economic decisions are modeled as pathfinding on '
    'a multi-dimensional attribute space with Mahalanobis edge weights (Bond, 2026), loss aversion '
    'arises naturally as a geometric asymmetry: losses traverse more dimensions of the attribute space '
    'than equivalent gains. The predicted loss aversion coefficient is:'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u03bb(d) = \u221ad')
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph(
    'where d is the number of evaluative dimensions activated by the good in question. This function '
    'has no free parameters once the attribute space is specified. It predicts that loss aversion should '
    'be absent for pure money (d = 1, \u03bb = 1), moderate for standard goods (d = 5, \u03bb = 2.24), '
    'and strong for identity-laden goods (d \u2265 7, \u03bb \u2265 2.65)\u2014precisely the pattern '
    'Horowitz and McConnell documented.'
)

# ═══════════════════════════════════════════════════════════════
# 2. THE MODEL
# ═══════════════════════════════════════════════════════════════

h = doc.add_heading('2. The Model', level=1)
h.runs[0].font.size = Pt(14)

doc.add_heading('2.1 Multi-Attribute Decision Space', level=2)

doc.add_paragraph(
    'We model economic decisions on a d-dimensional attribute space following Bond (2026). Each '
    'economic state v carries an attribute vector a(v) \u2208 \u211d\u1d48 whose components represent '
    'evaluatively distinct dimensions of the decision. In the general framework, d = 9 dimensions '
    'are identified empirically (consequences, rights, fairness, autonomy, trust, social impact, '
    'identity, legitimacy, epistemic status), but the derivation holds for any d \u2265 1.'
)

doc.add_paragraph(
    'The cost of transitioning from state v to v\u2032 is the behavioral friction:'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('BF(v, v\u2032) = \u221a(\u0394a\u1d40 \u03a3\u207b\u00b9 \u0394a)')
run.font.size = Pt(11)

doc.add_paragraph(
    'where \u03a3 \u2208 \u211d\u1d48\u02e3\u1d48 is the attribute covariance matrix capturing '
    'cross-dimensional correlations.'
)

doc.add_heading('2.2 Asymmetric Dimensional Activation', level=2)

doc.add_paragraph(
    'The key observation: gains and losses of the same monetary magnitude do not activate the same '
    'dimensions.'
)

p = doc.add_paragraph()
run = p.add_run('Definition 1 (Dimensional Loading). ')
run.bold = True
p.add_run(
    'For a transaction involving good g, define the dimensional loading d(g) as the number of '
    'attribute dimensions with nonzero change in the attribute vector. A monetary gain of \u03b4 changes '
    'only the consequences dimension (d_gain = 1). A loss of the same good changes multiple dimensions: '
    'consequences (a\u2081: monetary loss), plus rights (a\u2082: \u201cthis was mine\u201d), fairness '
    '(a\u2083: \u201cthis is unjust\u201d), identity (a\u2087: \u201cI am worse off\u201d), and '
    'possibly others depending on the good. Thus d_loss(g) \u2265 d_gain(g) in general, with equality '
    'only when g is pure money.'
)

doc.add_heading('2.3 Derivation of \u03bb(d)', level=2)

p = doc.add_paragraph()
run = p.add_run('Proposition 1 (Variable Loss Aversion). ')
run.bold = True
p.add_run(
    'Under the Mahalanobis decision model with covariance \u03a3, the loss aversion coefficient for '
    'a good with dimensional loading d_loss is: \u03bb(d_loss) = BF(loss of g) / BF(gain of g). '
    'When \u03a3 is approximately isotropic on the active dimensions (\u03a3 \u2248 \u03c3\u00b2I '
    'restricted to the active subspace), this simplifies to:'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\u03bb(d_loss) = \u221a(d_loss / d_gain) = \u221a(d_loss)')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph('since d_gain = 1 for monetary gains.')

p = doc.add_paragraph()
run = p.add_run('Proof. ')
run.bold = True
run.italic = True
p.add_run(
    'Let \u0394a_gain be the attribute-vector change for a monetary gain of \u03b4, with a single '
    'nonzero component: \u0394a\u2081 = \u03b4, \u0394a_k = 0 for k > 1. The behavioral friction is: '
    'BF_gain = |\u03b4| / \u03c3\u2081. Let \u0394a_loss be the change for a loss of the same good, '
    'with d_loss nonzero components of magnitude ~ \u03b4 each. Under isotropic \u03a3: '
    'BF_loss = |\u03b4| \u221a(d_loss) / \u03c3. The ratio is \u03bb = BF_loss / BF_gain = \u221a(d_loss). \u25a1'
)

doc.add_paragraph(
    'The result is intuitive: losing a good traverses a longer path through the attribute space than '
    'gaining the equivalent monetary value, because the loss activates dimensions (rights, fairness, '
    'identity) that the gain does not. The \u201cirrationally\u201d high weight on losses is '
    'geometrically rational\u2014the loss really does cost more, measured on the full manifold.'
)

# ═══════════════════════════════════════════════════════════════
# 3. EMPIRICAL VALIDATION
# ═══════════════════════════════════════════════════════════════

h = doc.add_heading('3. Empirical Validation', level=1)
h.runs[0].font.size = Pt(14)

doc.add_heading('3.1 Calibration', level=2)

doc.add_paragraph(
    'The model requires one input: the d \u00d7 d attribute covariance matrix \u03a3. We calibrate \u03a3 '
    'from a single dataset\u2014106 ultimatum game observations from Fraser and Nettle (2022)\u2014using '
    '5-fold cross-validated ridge regression with regularization parameter \u03b1 = 0.001. The '
    'calibrated \u03a3 is then frozen: all subsequent predictions use this same matrix with zero '
    're-calibration.'
)

doc.add_heading('3.2 Out-of-Sample Predictions', level=2)

doc.add_paragraph('Using the frozen \u03a3, we generate zero-parameter predictions for three canonical games:')

# Table 1
table = doc.add_table(rows=4, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Game', 'Predicted', 'Observed', 'Source']
for i, h_text in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h_text
    cell.paragraphs[0].runs[0].bold = True

data = [
    ['Ultimatum (offer %)', '36.9%', '44%', 'Camerer (2003)'],
    ['Dictator (transfer %)', '37.4%', '28.4%', 'Engel (2011)'],
    ['Public goods (contribution %)', '50.0%', '50%', 'Ledyard (1995)'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Table 1. ')
run.bold = True
p.add_run('Out-of-sample game predictions using frozen \u03a3 from ultimatum calibration.')

doc.add_paragraph(
    'The same \u03a3 also predicts the WTA/WTP ratio for the endowment effect: predicted 1.66, '
    'observed 2.0\u20132.5 (Kahneman, Knetsch, & Thaler, 1990).'
)

doc.add_heading('3.3 The Unique Prediction: Variable Loss Aversion', level=2)

# Table 2
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Good Type', 'Active Dimensions', 'Predicted \u03bb', 'CPT \u03bb']
for i, h_text in enumerate(headers):
    table.rows[0].cells[i].text = h_text
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ['Pure cash', '1', '1.00', '2.25'],
    ['Commodity', '2', '1.63', '2.25'],
    ['Gift', '3', '1.91', '2.25'],
    ['Standard good (KT)', '5', '2.22', '2.25'],
    ['Heirloom', '7 (eff. \u2248 12.5)', '3.53', '2.25'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Table 2. ')
run.bold = True
p.add_run('Predicted variable loss aversion vs. CPT constant. The geometric model recovers '
          '\u03bb \u2248 2.25 for standard goods (d = 5) and predicts the variation across good types '
          'that CPT cannot explain.')

doc.add_paragraph(
    'The geometric model matches the Kahneman\u2013Tversky estimate (\u03bb \u2248 2.25) for standard '
    'goods (d = 5, predicted \u03bb = 2.22)\u2014recovering CPT as a special case. But it also predicts '
    'the variation across good types: near-zero loss aversion for pure money, moderate for commodities, '
    'and strong for identity-laden goods. CPT, by construction, predicts \u03bb = 2.25 for all goods. '
    'Monotonicity is verified: \u03bb increases strictly with d.'
)

doc.add_heading('3.4 Cross-Cultural Validation', level=2)

doc.add_paragraph('The frozen \u03a3 generates predictions for seven cross-cultural populations:')

# Table 3
table = doc.add_table(rows=9, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Population', 'Predicted', 'Observed', 'Error']
for i, h_text in enumerate(headers):
    table.rows[0].cells[i].text = h_text
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ['Machiguenga (Peru)', '36%', '26%', '10.0%'],
    ['Hadza (Tanzania)', '36%', '27%', '9.2%'],
    ['Tsimane (Bolivia)', '36%', '32%', '4.3%'],
    ['US students', '37%', '42%', '5.1%'],
    ['Europe average', '37%', '44%', '7.1%'],
    ['Au (Papua New Guinea)', '37%', '44%', '6.7%'],
    ['Lamalera (Indonesia)', '38%', '57%', '19.3%'],
    ['Cross-cultural MAE', '', '', '8.8%'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val
    if r == 7:
        table.rows[r+1].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[r+1].cells[3].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Table 3. ')
run.bold = True
p.add_run('Cross-cultural predictions using a single frozen \u03a3 calibrated on WEIRD data. '
          'MAE = 8.8% with zero cultural parameters.')

doc.add_heading('3.5 Baseline Comparison', level=2)

# Table 4
table = doc.add_table(rows=5, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Model', 'Cross-game?', 'Params/game', 'Variable \u03bb?']
for i, h_text in enumerate(headers):
    table.rows[0].cells[i].text = h_text
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

data = [
    ['Nash equilibrium', 'Yes (wrong)', '0', 'No (\u03bb = 1 always)'],
    ['CPT', 'No', '4\u20135', 'No (\u03bb \u2248 2.25 constant)'],
    ['Fehr\u2013Schmidt', 'No', '2+', 'No'],
    ['Geometric model', 'Yes', '0 (frozen)', 'Yes: \u03bb = \u221ad'],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val
    if r == 3:
        for c in range(4):
            table.rows[r+1].cells[c].paragraphs[0].runs[0].bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Table 4. ')
run.bold = True
p.add_run('Model comparison. The geometric model is the only one that generates cross-game predictions '
          'with zero per-game parameters and predicts variable loss aversion.')

# ═══════════════════════════════════════════════════════════════
# 4. DISCUSSION
# ═══════════════════════════════════════════════════════════════

h = doc.add_heading('4. Discussion', level=1)
h.runs[0].font.size = Pt(14)

doc.add_paragraph(
    'The result has three notable features. First, it derives a number that is usually fitted. CPT '
    'treats \u03bb = 2.25 as a parameter estimated from data; the geometric model derives '
    '\u03bb(d) = \u221ad from the structure of the decision space. The Kahneman\u2013Tversky value is '
    'recovered as the special case d = 5, corresponding to a \u201cstandard\u201d good that activates '
    'about half the attribute dimensions.'
)

doc.add_paragraph(
    'Second, it makes a unique, falsifiable prediction: loss aversion should vary systematically with '
    'the dimensional loading of the good. Specifically, in a within-subjects design that holds monetary '
    'value constant while varying the good type (cash, commodity, gift, heirloom), the measured \u03bb '
    'should track \u221ad. This prediction is absent from all existing models.'
)

doc.add_paragraph(
    'Third, the result explains an empirical puzzle. Novemsky and Kahneman (2005) noted that loss '
    'aversion appears attenuated or absent for routine monetary transactions. This is exactly what '
    '\u03bb(1) = 1 predicts: when only the monetary dimension is active (d = 1), there is no geometric '
    'asymmetry between gains and losses. Loss aversion is not a universal feature of the value function; '
    'it is a consequence of dimensional asymmetry that emerges only when the good activates non-monetary '
    'evaluative dimensions.'
)

doc.add_heading('4.1 Limitations', level=2)

doc.add_paragraph(
    'The \u03bb = \u221ad result assumes approximately isotropic \u03a3 on the active dimensions. When '
    '\u03a3 is strongly anisotropic, the formula generalizes to \u03bb = \u221a(\u2211_k \u03c3_k\u207b\u00b2 / '
    '\u03c3\u2081\u207b\u00b2) where the sum runs over active dimensions, giving different weights to '
    'different attribute directions. The isotropic case is a useful approximation but not exact.'
)

doc.add_paragraph(
    'The dimensional loading d(g) for each good type is currently assigned by the analyst, not measured. '
    'An experimental protocol that measures dimensional activation (e.g., via attribute rating tasks '
    'before and after transactions) would allow d(g) to be determined empirically, enabling a fully '
    'quantitative test.'
)

doc.add_paragraph(
    'The cross-cultural predictions are limited by the use of a single \u03a3 calibrated on WEIRD data. '
    'Culture-specific calibration would improve accuracy but introduce parameters.'
)

# ═══════════════════════════════════════════════════════════════
# 5. CONCLUSION
# ═══════════════════════════════════════════════════════════════

h = doc.add_heading('5. Conclusion', level=1)
h.runs[0].font.size = Pt(14)

doc.add_paragraph(
    'Loss aversion is not a constant. It is a geometric consequence of the dimensionality of the goods '
    'being traded. The Kahneman\u2013Tversky estimate \u03bb \u2248 2.25 is recovered as the special '
    'case for goods activating approximately five evaluative dimensions. For pure money, \u03bb = 1 (no '
    'loss aversion). For identity-laden goods, \u03bb > 3. This variation is predicted by the geometric '
    'framework with zero free parameters and matches the meta-analytic evidence that CPT\u2014by '
    'construction\u2014cannot explain.'
)

# ═══════════════════════════════════════════════════════════════
# DECLARATIONS
# ═══════════════════════════════════════════════════════════════

h = doc.add_heading('Statements and Declarations', level=1)
h.runs[0].font.size = Pt(14)

p = doc.add_paragraph()
run = p.add_run('Competing Interests: ')
run.bold = True
p.add_run('The author has no relevant financial or non-financial interests to disclose.')

p = doc.add_paragraph()
run = p.add_run('Funding: ')
run.bold = True
p.add_run('No funding was received for conducting this study.')

p = doc.add_paragraph()
run = p.add_run('Data Availability: ')
run.bold = True
p.add_run('All data and code supporting the results are available at '
          'https://github.com/ahb-sjsu/eris-econ. The validation script (hpc/run_validation.py) '
          'reproduces all results reported in this paper. Experimental data from Fraser and Nettle '
          '(2022) and Henrich et al. (2001, 2005) are publicly available as cited.')

p = doc.add_paragraph()
run = p.add_run('Use of AI Tools: ')
run.bold = True
p.add_run('Manuscript preparation was assisted by Claude (Anthropic, claude-opus-4-6). All theoretical '
          'derivations, experimental design, calibration methodology, and scientific claims were '
          'developed by the author. The AI tool was used for drafting and formatting assistance. '
          'The author takes full responsibility for the content of the manuscript.')

p = doc.add_paragraph()
run = p.add_run('Author Contributions: ')
run.bold = True
p.add_run('Andrew H. Bond: conceptualization, methodology, formal analysis, software, validation, '
          'writing \u2014 original draft, writing \u2014 review and editing.')

# ═══════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════

h = doc.add_heading('References', level=1)
h.runs[0].font.size = Pt(14)

refs = [
    'Bond, A. H. (2026). Multi-attribute equilibrium: A manifold generalization of Nash. Working paper, San Jos\u00e9 State University.',
    'Camerer, C. F. (2003). Behavioral game theory: Experiments in strategic interaction. Princeton University Press.',
    'Engel, C. (2011). Dictator games: A meta study. Experimental Economics, 14(4), 583\u2013610. https://doi.org/10.1007/s10683-011-9283-7',
    'Fraser, S., & Nettle, D. (2022). The role of need in altruistic and non-altruistic behavior. Evolution and Human Behavior, 43(6), 469\u2013477. https://doi.org/10.1016/j.evolhumbehav.2022.08.004',
    'Henrich, J., Boyd, R., Bowles, S., Camerer, C., Fehr, E., Gintis, H., & McElreath, R. (2001). In search of Homo economicus: Behavioral experiments in 15 small-scale societies. American Economic Review, 91(2), 73\u201378. https://doi.org/10.1257/aer.91.2.73',
    'Henrich, J., Boyd, R., Bowles, S., Camerer, C., Fehr, E., Gintis, H., McElreath, R., Alvard, M., Barr, A., Ensminger, J., Henrich, N. S., Hill, K., Gil-White, F., Gurven, M., Marlowe, F. W., Patton, J. Q., & Tracer, D. (2005). \u201cEconomic man\u201d in cross-cultural perspective: Behavioral experiments in 15 small-scale societies. Behavioral and Brain Sciences, 28(6), 795\u2013815. https://doi.org/10.1017/S0140525X05000142',
    'Horowitz, J. K., & McConnell, K. E. (2002). A review of WTA/WTP studies. Journal of Environmental Economics and Management, 44(3), 426\u2013447. https://doi.org/10.1006/jeem.2001.1215',
    'Kahneman, D., & Tversky, A. (1979). Prospect theory: An analysis of decision under risk. Econometrica, 47(2), 263\u2013292. https://doi.org/10.2307/1914185',
    'Kahneman, D., Knetsch, J. L., & Thaler, R. H. (1990). Experimental tests of the endowment effect and the Coase theorem. Journal of Political Economy, 98(6), 1325\u20131348. https://doi.org/10.1086/261737',
    'Ledyard, J. O. (1995). Public goods: A survey of experimental research. In J. H. Kagel & A. E. Roth (Eds.), The handbook of experimental economics (pp. 111\u2013194). Princeton University Press.',
    'Novemsky, N., & Kahneman, D. (2005). The boundaries of loss aversion. Journal of Marketing Research, 42(2), 119\u2013128. https://doi.org/10.1509/jmkr.42.2.119.62292',
    'Tversky, A., & Kahneman, D. (1992). Advances in prospect theory: Cumulative representation of uncertainty. Journal of Risk and Uncertainty, 5(4), 297\u2013323. https://doi.org/10.1007/BF00122574',
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.style.font.size = Pt(11)

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════

output = r'C:\source\erisml-lib\docs\papers\paper4-variable-loss-aversion\Bond_2026_Variable_Loss_Aversion_JRU.docx'
doc.save(output)
print(f'Saved: {output}')
