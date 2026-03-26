"""Add figures to the existing JRU .docx."""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

DOCX = r'C:\source\erisml-lib\docs\papers\paper4-variable-loss-aversion\Bond_2026_Variable_Loss_Aversion_JRU.docx'
FIGS = Path(r'C:\source\erisml-lib\docs\papers\paper4-variable-loss-aversion\figures')

doc = Document(DOCX)

# Find insertion points by searching for specific text
def find_para(doc, text_fragment):
    for i, p in enumerate(doc.paragraphs):
        if text_fragment in p.text:
            return i
    return None

# Figure 1: After "precisely the pattern Horowitz and McConnell documented."
idx = find_para(doc, 'precisely the pattern Horowitz and McConnell documented')
if idx:
    # Insert after this paragraph — we add to the end of section 1
    # Find the "2. The Model" heading and insert before it
    model_idx = find_para(doc, '2. The Model')
    if model_idx:
        # Add figure before "2. The Model"
        p = doc.paragraphs[model_idx]
        # We need to insert before this element
        el = p._element

        # Create figure paragraph
        from docx.oxml.ns import qn
        from lxml import etree
        import copy

        # Simpler: just add at end and note positions
        pass

# Since inserting at arbitrary positions in python-docx is complex,
# let's rebuild the doc with figures in the right places
from docx import Document as Doc2
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read existing doc
doc = Document(DOCX)

# Find key paragraphs and insert figures after them
fig_insertions = {
    'precisely the pattern Horowitz and McConnell documented': (
        FIGS / 'fig1_lambda_vs_dimensions.png',
        'Fig 1. Predicted loss aversion \u03bb = \u221ad (solid curve) versus the CPT constant \u03bb = 2.25 (dashed). '
        'Data points show model predictions for goods with different dimensional loadings. The geometric model '
        'captures the variation documented by Horowitz and McConnell (2002); CPT cannot'
    ),
    'the loss really does cost more, measured on the full manifold': (
        FIGS / 'fig2_dimensional_activation.png',
        'Fig 2. Dimensional activation for gains versus losses. (a) A monetary gain activates only the consequences '
        'dimension (d = 1). (b) Loss of an heirloom activates seven dimensions. The Mahalanobis path length is '
        '\u221a7 times longer for the loss, yielding \u03bb = 2.65'
    ),
    'Culture-specific calibration would improve accuracy but introduce parameters': (
        FIGS / 'fig3_cross_cultural.png',
        'Fig 3. Cross-cultural ultimatum game predictions (frozen \u03a3, zero cultural parameters) versus '
        'observed offers across seven populations. Numbers above bars show absolute error. MAE = 8.8%'
    ),
}

# Also add fig 4 after the proof
fig4_insert = {
    'since d_gain = 1 for monetary gains': (
        FIGS / 'fig4_geometric_intuition.png',
        'Fig 4. Geometric intuition for variable loss aversion. A gain traverses only the monetary axis '
        '(short path). A loss of the same good traverses both the monetary and fairness axes (longer diagonal path). '
        'The ratio of path lengths gives \u03bb = \u221a2 for a good with d = 2'
    ),
}
fig_insertions.update(fig4_insert)

# Process: find each target paragraph, add figure paragraph after it
for target_text, (fig_path, caption) in fig_insertions.items():
    idx = find_para(doc, target_text)
    if idx is None:
        print(f'WARNING: Could not find "{target_text[:50]}..."')
        continue

    # python-docx doesn't support easy insertion, so we'll add figures
    # by appending after the target paragraph's XML element
    target_elem = doc.paragraphs[idx]._element

    # Create image paragraph
    from docx.oxml import OxmlElement

    # Add blank line
    new_p = OxmlElement('w:p')
    target_elem.addnext(new_p)

    # Add caption paragraph after the blank
    caption_p = OxmlElement('w:p')
    new_p.addnext(caption_p)

    # Actually, this low-level approach is fragile. Let me use a simpler method:
    # Add all figures at the end, with clear "INSERT FIGURE X HERE" markers
    pass

# Simpler approach: just add all figures at the end as an appendix,
# with clear figure numbers that match the text references.
# OR: rebuild the document from scratch with figures inline.

# Let's take the pragmatic approach: add figures inline by finding
# paragraphs and using the run.add_picture method on a new paragraph

print("Adding figures to document...")

for target_text, (fig_path, caption_text) in fig_insertions.items():
    idx = find_para(doc, target_text)
    if idx is None:
        print(f'  SKIP: {target_text[:40]}...')
        continue

    # Get the XML element of the target paragraph
    target = doc.paragraphs[idx]._element
    parent = target.getparent()

    # Create a new paragraph for the figure
    from docx.oxml.ns import qn

    # Figure paragraph
    fig_p = doc.add_paragraph()
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_p.add_run()
    run.add_picture(str(fig_path), width=Inches(5.0))

    # Caption paragraph
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(caption_text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)

    # Blank line
    blank_p = doc.add_paragraph()

    # Move these paragraphs to right after the target
    # (they were added at the end, need to relocate)
    for p_to_move in [fig_p, cap_p, blank_p]:
        elem = p_to_move._element
        parent_body = elem.getparent()
        parent_body.remove(elem)
        target.addnext(elem)
        target = elem  # chain: each new element goes after the previous

    print(f'  Added: {fig_path.name} after "{target_text[:40]}..."')

doc.save(DOCX)
print(f'\nSaved: {DOCX}')
