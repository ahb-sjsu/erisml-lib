#!/usr/bin/env python3
"""
Build Geometric Ethics v1.18 from v1.16.
Adds Part VI: Domain Applications with nine chapters:
  - Chapter 20: Geometric Economics (BGE, Nash nesting)
  - Chapter 21: Geometric Clinical Ethics
  - Chapter 22: Geometric Jurisprudence
  - Chapter 23: Geometric Finance
  - Chapter 24: Geometric Theology
  - Chapter 25: Geometric Environmental Ethics
  - Chapter 26: Geometric AI Ethics
  - Chapter 27: Geometric Bioethics
  - Chapter 28: Geometric Military Ethics
Renumbers old Part VI -> Part VII, old Ch 20 -> Ch 29, Ch 21 -> Ch 30.
Updates cross-references, Arc of the Book, Arc of the Argument, version.
"""

import copy
import re
import sys
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

SRC = Path(r"C:\Users\abptl\Documents\Geometric Ethics - The Mathematical Structure of Moral Reasoning - Bond - v1.16 - Mar 2026.docx")
DST = Path(r"C:\Users\abptl\Documents\Geometric Ethics - The Mathematical Structure of Moral Reasoning - Bond - v1.18 - Mar 2026.docx")

print("Loading v1.16...")
doc = Document(str(SRC))

# ============================================================
# STEP 1: Find insertion point (just before "Part VI: Horizons")
# ============================================================
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.style.name == 'Heading 1' and 'Part VI' in p.text and 'Horizon' in p.text:
        insert_idx = i
        break

if insert_idx is None:
    print("ERROR: Could not find 'Part VI: Horizons'")
    sys.exit(1)

print(f"Found Part VI: Horizons at paragraph {insert_idx}")

# ============================================================
# STEP 2: Renumber old Part VI -> Part VII, Ch 20 -> 29, Ch 21 -> 30
# ============================================================
XREF_MAP = {
    'Chapter 20': 'Chapter 29',
    'Chapter 21': 'Chapter 30',
    'Ch. 20': 'Ch. 29',
    'Ch. 21': 'Ch. 30',
    'Ch 20': 'Ch 29',
    'Ch 21': 'Ch 30',
    '\u00a720.': '\u00a729.',
    '\u00a721.': '\u00a730.',
    'Section 20.': 'Section 29.',
    'Section 21.': 'Section 30.',
    'Part VI': 'Part VII',
}

def renumber_text(text):
    """Apply all renumbering to a text string."""
    result = text
    for old, new in XREF_MAP.items():
        result = result.replace(old, new)
    return result

def renumber_runs(paragraph):
    """Renumber cross-references in a paragraph while preserving formatting."""
    for run in paragraph.runs:
        new_text = renumber_text(run.text)
        if new_text != run.text:
            run.text = new_text

print("Renumbering Part VI -> Part VII, Ch 20 -> 29, Ch 21 -> 30...")
renumber_count = 0
for p in doc.paragraphs:
    old_text = p.text
    renumber_runs(p)
    if p.text != old_text:
        renumber_count += 1

print(f"  Renumbered text in {renumber_count} paragraphs")

# Renumber section numbers (20.x -> 29.x, 21.x -> 30.x) in ALL paragraphs.
# This catches headings, body text, index entries, and appendix references.
# Safe to apply globally because this runs BEFORE new Ch 20-28 content is inserted.
section_renumber_count = 0
for p in doc.paragraphs:
    for run in p.runs:
        new_text = re.sub(r'\b20\.(\d)', r'29.\1', run.text)
        new_text = re.sub(r'\b21\.(\d)', r'30.\1', new_text)
        if new_text != run.text:
            run.text = new_text
            section_renumber_count += 1

print(f"  Renumbered {section_renumber_count} section references (20.x -> 29.x, 21.x -> 30.x)")

# ============================================================
# STEP 3: Update version number in header/title
# ============================================================
for p in doc.paragraphs:
    if 'v 1.16' in p.text or 'v1.16' in p.text:
        for run in p.runs:
            run.text = run.text.replace('v 1.16', 'v 1.18').replace('v1.16', 'v1.18')

# Update headers/footers
for section in doc.sections:
    for header in [section.header, section.first_page_header]:
        if header:
            for p in header.paragraphs:
                for run in p.runs:
                    run.text = run.text.replace('v 1.16', 'v 1.18').replace('v1.16', 'v1.18')

print("Updated version to v1.18")

# ============================================================
# STEP 3a: Patch Preface chapter/part counts and audience section
# ============================================================
for p in doc.paragraphs:
    for run in p.runs:
        if 'twenty-one chapters across six parts' in run.text:
            run.text = run.text.replace('twenty-one chapters across six parts',
                                        'thirty chapters across seven parts')
            print("  Patched Preface: twenty-one chapters -> thirty chapters")
        if 'twenty-one chapters' in run.text and 'six parts' in run.text:
            run.text = run.text.replace('twenty-one chapters', 'thirty chapters')
            run.text = run.text.replace('six parts', 'seven parts')

# Add Part VI audience guidance after the policy makers paragraph
for i, p in enumerate(doc.paragraphs):
    if 'Policy makers and governance professionals will find' in p.text:
        # Insert after this paragraph
        from lxml import etree as et2
        from docx.oxml.ns import qn as qn2
        target_aud = doc.paragraphs[i]._element
        new_aud = copy.deepcopy(target_aud)
        for child in list(new_aud):
            new_aud.remove(child)
        pPr_aud = et2.SubElement(new_aud, qn2('w:pPr'))
        pStyle_aud = et2.SubElement(pPr_aud, qn2('w:pStyle'))
        pStyle_aud.set(qn2('w:val'), 'BodyText')
        run_aud = et2.SubElement(new_aud, qn2('w:r'))
        t_aud = et2.SubElement(run_aud, qn2('w:t'))
        t_aud.text = (
            'Domain specialists \u2014 economists, clinicians, lawyers, financial professionals, theologians, '
            'environmental scientists, AI researchers, bioethicists, and military ethicists \u2014 will find '
            'direct application in Part VI: Domain Applications (Chapters 20\u201328), where the framework is '
            'applied to nine established domains, each with worked examples, formal theorems, and falsifiable '
            'predictions distinguishing the geometric approach from existing domain-specific theories.'
        )
        t_aud.set(qn2('xml:space'), 'preserve')
        target_aud.addnext(new_aud)
        print("  Inserted domain specialist audience paragraph in Preface")
        break

# ============================================================
# STEP 3b: Patch "Arc of the Book" and "Arc of the Argument"
# ============================================================
from lxml import etree
from docx.oxml.ns import qn

# Insert Part VI description in the Preface's "Arc of the Book" section
for i, p in enumerate(doc.paragraphs):
    if 'Part VII: Horizons looks forward' in p.text:
        print(f"Found Part VII reference in Arc of the Book at paragraph {i}")
        target_arc = p._element
        new_para_arc = copy.deepcopy(target_arc)
        for child in list(new_para_arc):
            new_para_arc.remove(child)
        pPr_arc = etree.SubElement(new_para_arc, qn('w:pPr'))
        pStyle_arc = etree.SubElement(pPr_arc, qn('w:pStyle'))
        pStyle_arc.set(qn('w:val'), 'BodyText')
        run_arc = etree.SubElement(new_para_arc, qn('w:r'))
        t_arc = etree.SubElement(run_arc, qn('w:t'))
        t_arc.text = (
            'Part VI: Domain Applications demonstrates that the framework is not confined to abstract ethics or AI alignment. '
            'Chapter 20 applies the framework to economics, constructing the Bond Geodesic Equilibrium \u2014 a generalization of Nash equilibrium in which agents optimize on the full manifold \u2014 and proving that Nash equilibrium is the scalar projection of BGE. '
            'Chapter 21 applies the framework to clinical medicine, showing that the QALY is a scalar projection that destroys eight dimensions of clinically relevant information, that moral injury is cumulative manifold damage distinct from burnout, and that informed consent is a gauge-invariance condition. '
            'Chapter 22 applies the framework to law, constructing the judicial complex, deriving the octahedral gauge group D\u2084 \u22cb D\u2084 from the full Hohfeldian octad, and recasting constitutional review as path homology preservation. '
            'Chapter 23 applies the framework to financial markets, identifying risk as manifold curvature, the Flash Crash as dimensional collapse, and the implied volatility surface as the shadow of higher-dimensional pricing. '
            'Chapter 24 applies the framework to theology, showing that the moral manifold is cross-religiously invariant, that the Genesis Fall is an epistemic impossibility theorem, and that the Euthyphro dilemma is a question about gauge invariance resolved by the empirical data. '
            'Chapter 25 applies the framework to environmental ethics, formalizing intergenerational pathfinding, the discount rate as dimensional collapse, the tragedy of the commons as multi-agent manifold failure, and species extinction as irreversible boundary crossing. '
            'Chapter 26 applies the framework to AI ethics, diagnosing the alignment problem as scalar irrecoverability, algorithmic bias as gauge-invariance violation, and the paperclip maximizer as total dimensional collapse. '
            'Chapter 27 applies the framework to population-level bioethics, formalizing CRISPR germline editing as irreversible manifold modification, the enhancement equity problem, and research consent as a double gauge condition. '
            'Chapter 28 applies the framework to military ethics, recasting just war theory as manifold entry conditions, proportionality as multi-dimensional cost-benefit, and autonomous weapons as dimensional collapse in lethal domains.'
        )
        t_arc.set(qn('xml:space'), 'preserve')
        target_arc.addprevious(new_para_arc)
        print("  Inserted Part VI description in Arc of the Book")
        break

# Insert item 8 in the Conclusion's "Arc of the Argument" section
for i, p in enumerate(doc.paragraphs):
    if '7. The framework applies to AI systems' in p.text:
        print(f"Found Arc of Argument item 7 at paragraph {i}")
        target_arg = doc.paragraphs[i+1]._element
        new_para_arg = copy.deepcopy(doc.paragraphs[i]._element)
        for child in list(new_para_arg):
            new_para_arg.remove(child)
        pPr_arg = etree.SubElement(new_para_arg, qn('w:pPr'))
        pStyle_arg = etree.SubElement(pPr_arg, qn('w:pStyle'))
        pStyle_arg.set(qn('w:val'), 'BodyText')
        run_arg = etree.SubElement(new_para_arg, qn('w:r'))
        t_arg = etree.SubElement(run_arg, qn('w:t'))
        t_arg.text = (
            '8. The framework applies across domains (Chapters 20\u201328). Nine domains \u2014 economics, clinical medicine, law, finance, theology, environmental ethics, AI ethics, bioethics, and military ethics \u2014 each instantiate the moral manifold with domain-specific dimensions, boundaries, and metrics. '
            'The Bond Geodesic Equilibrium subsumes Nash equilibrium. The QALY Irrecoverability Theorem shows scalar clinical measures destroy eight dimensions of information. '
            'Topological constitutionality recasts judicial review as path homology preservation. Risk is manifold curvature. The moral manifold is cross-religiously invariant. '
            'The discount rate is dimensional collapse destroying intergenerational information. The alignment problem is scalar irrecoverability applied to reward functions. '
            'Germline editing is irreversible manifold modification. Just war proportionality is multi-dimensional, not scalar. '
            'Each domain produces falsifiable predictions that distinguish the geometric approach from existing domain-specific theories.'
        )
        t_arg.set(qn('xml:space'), 'preserve')
        target_arg.addprevious(new_para_arg)
        print("  Inserted item 8 in Arc of the Argument")
        break

# ============================================================
# STEP 4: Build the new chapter content
# ============================================================

# Import chapter content from separate files
sys.path.insert(0, str(Path(__file__).parent))
from ch20_economics import CH20_ECONOMICS
from ch21_clinical import CH21_CLINICAL
from ch22_jurisprudence import CH22_JURISPRUDENCE
from ch23_finance import CH23_FINANCE
from ch24_theology import CH24_THEOLOGY
from ch25_environmental import CH25_ENVIRONMENTAL
from ch26_ai import CH26_AI
from ch27_bioethics import CH27_BIOETHICS
from ch28_military import CH28_MILITARY

# Part VI intro
PART_INTRO = [
    ('Heading 1', 'Part VI: Domain Applications'),
    ('Body Text', 'The preceding five parts developed the mathematical framework of Geometric Ethics: the moral manifold, tensor hierarchy, dynamics, symmetry, conservation laws, and implementation architecture. This part demonstrates that the framework is not confined to abstract ethics or AI alignment. The same mathematical structures \u2014 pathfinding on stratified manifolds, gauge invariance, Noether conservation, tensorial contraction \u2014 apply directly to established domains with their own formal traditions: economics, clinical medicine, law, finance, theology, environmental policy, artificial intelligence, bioethics, and military ethics.'),
    ('Body Text', 'Each chapter in this part takes a domain that has struggled with the limitations of scalar models, shows how the geometric framework resolves specific longstanding puzzles, and identifies falsifiable predictions that distinguish the geometric approach from existing domain-specific theories. The chapters can be read independently, but they share a common architecture: domain-specific instantiation of the moral manifold, identification of the relevant dimensions, construction of the domain geodesic, derivation of results inaccessible from scalar models, worked examples applying the framework to real-world cases, and falsifiable predictions. The nine domains span the full range of human moral decision-making, from individual clinical encounters to global climate policy, from ancient just war doctrine to cutting-edge AI alignment.'),
    ('Body Text', ''),
    ('Heading 2', 'Methodological Notes for Domain Applications'),
    ('Body Text', 'Before proceeding, three methodological concerns that pervade all nine domain chapters deserve explicit treatment: the grounding problem, computational tractability, and the calibration of the covariance matrix.'),
    ('Body Text', 'The Grounding Tensor and the Is\u2013Ought Gap. Each domain chapter constructs a decision complex whose vertices carry nine-dimensional attribute vectors. These vectors are computed from observable data by a grounding function \u03a8 that maps physical observables (pixels, text, sensor readings, market data, medical records) to moral-dimensional scores. The philosophical challenge is immediate: \u03a8 is where the is\u2013ought gap lives. How does one derive a d_7 (virtue/identity) score from raw data? The framework does not claim to have dissolved the is\u2013ought gap. Rather, it has localized it. In conventional AI systems and decision-support tools, the mapping from observation to moral evaluation is implicit, distributed across training data, loss functions, and architectural choices \u2014 invisible, unauditable, and ungovernable. The geometric framework makes \u03a8 an explicit, inspectable, version-controlled software layer with defined inputs, outputs, and calibration procedures. Chapter 17 demonstrated that \u03a8 can be empirically calibrated: linear probes trained on the cross-lingual validation corpus achieved F_1 = 0.74\u20130.91 across the nine dimensions, with independent replication (Thiele, 2026). The grounding function is imperfect \u2014 all measurement instruments are \u2014 but it is explicit, testable, and improvable. Making the is\u2013ought gap a governable engineering interface rather than an invisible black-box assumption is itself a substantial advance.'),
    ('Body Text', 'Computational Tractability. Exact geodesic computation on the full nine-dimensional moral manifold is computationally intractable (Theorem 11.2). The domain chapters rely on two sources of tractability. First, A* search with admissible heuristics provides polynomial-time approximate solutions with provably bounded suboptimality (Chapter 11). Domain-specific heuristics \u2014 moral rules in ethics, clinical guidelines in medicine, legal doctrines in law, trading rules in finance, ROE in military contexts \u2014 are the domain-specific instantiations of h(n) that make real-time pathfinding feasible. Second, in practice, not all nine dimensions are equally active in every decision context. A routine financial transaction may activate primarily d_1 (return) and d_2 (contractual obligation) with minimal activation of d_7 (identity) or d_9 (epistemic status). This dimensional sparsity reduces the effective dimensionality of the computation. For AI implementation, the DEME architecture (Chapter 19) employs Tucker decomposition and tensor-train formats to compress the rank-6 tensor operations, achieving sub-second inference on standard hardware for typical decision contexts. High-curvature regions (morally fraught decisions where many dimensions are active) require more computation \u2014 which is itself a desirable property: the system spends more time on hard moral decisions, mirroring human moral deliberation.'),
    ('Body Text', 'Calibration of the Covariance Matrix. Every domain chapter relies on a covariance matrix \u03a3 that encodes the statistical relationships among the nine moral dimensions within that domain. Estimating a 9\u00d79 positive-definite matrix (up to 45 free parameters in the symmetric case) from behavioral data raises legitimate concerns about identifiability, particularly for latent dimensions such as d_7 (virtue/identity) and d_9 (epistemic status) that are not directly observable. The framework addresses this challenge through three complementary strategies. First, the dimensional scores are not estimated from raw behavior but from the calibrated probes of Chapter 17, which map observable text and behavioral indicators to dimensional scores with known accuracy. The covariance matrix is then estimated from probe-scored data, not from unstructured observables. Second, structured experimental designs \u2014 discrete choice experiments, factorial vignette studies, and conjoint analyses \u2014 can orthogonalize the dimensional contributions, enabling identification of the covariance parameters via structural equation modeling (SEM) or maximum likelihood estimation (MLE). Third, the framework generates falsifiable predictions (six per domain chapter) that provide external validation: if the estimated \u03a3 produces predictions that fail empirically, the matrix is miscalibrated and must be re-estimated. The covariance matrix is not assumed; it is empirically estimated, cross-validated, and falsifiably constrained.'),
    ('Body Text', ''),
]

# Assemble all domain chapters
ALL_CHAPTERS = (PART_INTRO + CH20_ECONOMICS + CH21_CLINICAL + CH22_JURISPRUDENCE +
                CH23_FINANCE + CH24_THEOLOGY + CH25_ENVIRONMENTAL + CH26_AI +
                CH27_BIOETHICS + CH28_MILITARY)

# ============================================================
# STEP 5: Insert all content before Part VII (was Part VI)
# ============================================================
print(f"Inserting new Part VI: Domain Applications ({len(ALL_CHAPTERS)} paragraphs) before paragraph {insert_idx}...")

target_element = doc.paragraphs[insert_idx]._element

for style, text in ALL_CHAPTERS:
    new_para = etree.SubElement(target_element.getparent(), qn('w:p'))
    new_para.getparent().remove(new_para)

    new_para = copy.deepcopy(doc.paragraphs[0]._element)
    for child in list(new_para):
        new_para.remove(child)

    pPr = etree.SubElement(new_para, qn('w:pPr'))
    pStyle = etree.SubElement(pPr, qn('w:pStyle'))
    style_val = style.replace(' ', '')
    pStyle.set(qn('w:val'), style_val)

    run_elem = etree.SubElement(new_para, qn('w:r'))
    t_elem = etree.SubElement(run_elem, qn('w:t'))
    t_elem.text = text
    t_elem.set(qn('xml:space'), 'preserve')

    target_element.addprevious(new_para)

print(f"Inserted {len(ALL_CHAPTERS)} paragraphs")

# ============================================================
# STEP 6: Save
# ============================================================
print(f"Saving to {DST}...")
doc.save(str(DST))
print("Done! v1.18 saved.")
print()
print(f"Total domain application paragraphs: {len(ALL_CHAPTERS)}")
print(f"  Part Intro: {len(PART_INTRO)}")
print(f"  Ch 20 Economics: {len(CH20_ECONOMICS)}")
print(f"  Ch 21 Clinical: {len(CH21_CLINICAL)}")
print(f"  Ch 22 Jurisprudence: {len(CH22_JURISPRUDENCE)}")
print(f"  Ch 23 Finance: {len(CH23_FINANCE)}")
print(f"  Ch 24 Theology: {len(CH24_THEOLOGY)}")
print(f"  Ch 25 Environmental: {len(CH25_ENVIRONMENTAL)}")
print(f"  Ch 26 AI Ethics: {len(CH26_AI)}")
print(f"  Ch 27 Bioethics: {len(CH27_BIOETHICS)}")
print(f"  Ch 28 Military: {len(CH28_MILITARY)}")
print()
print("=== MANUAL STEPS REQUIRED ===")
print("1. Open v1.18 in Word and regenerate the Table of Contents")
print("2. Update the title page version number if not caught by script")
print("3. Review cross-references: all old 'Chapter 20/21' refs now point to 'Chapter 29/30'")
print("4. Add new chapters to the Index")
print("5. Update 'Core Objects at a Glance' table to include new domain constructs")
print("6. Update 'Key Results at a Glance' table to include new theorems")
print("7. Update Preface to mention Part VI: Domain Applications")
