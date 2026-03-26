"""Apply reviewer corrections to create v1.21 of the Geometric Ethics manuscript.

Fixes 5 issues identified by reviewer:
1. Killing field orthogonality overclaim (Lemma 9.2)
2. Convex cone claim (Theorem 9.1 / admissible metrics)
3. CHSH SU(2) exclusion overclaim (Corollary 12.1)
4. Conservation of harm rhetoric (modeling vs discovery)
5. A* intractability formalization (add complexity reduction)
"""

import docx
import copy
import sys
import os

SRC = r"C:\Users\abptl\Documents\Geometric Ethics - The Mathematical Structure of Moral Reasoning - Bond - v1.20 - Mar 2026.docx"
DST = r"C:\Users\abptl\Documents\Geometric Ethics - The Mathematical Structure of Moral Reasoning - Bond - v1.21 - Mar 2026.docx"

print(f"Reading: {SRC}")
doc = docx.Document(SRC)

fixes_applied = 0

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue

    # ═══════════════════════════════════════════════════════════════
    # FIX 1: Killing field / orthogonality (paragraph ~3024-3025)
    # ═══════════════════════════════════════════════════════════════
    if "Lemma 9.2" in text and "Killing vector field" in text and "g_{7" in text:
        # Find the run containing the overclaim
        for run in p.runs:
            if "forces g_{7μ} = 0 for μ ≠ 7" in run.text:
                run.text = run.text.replace(
                    "forces g_{7μ} = 0 for μ ≠ 7",
                    "forces ∂₇g_{μν} = 0 for all μ, ν (the metric is independent of x⁷)"
                )
                fixes_applied += 1
                print(f"  FIX 1a applied at paragraph {i}: Lemma statement")
            elif "The Killing equation ensures that g_{7μ} = 0" in run.text:
                run.text = run.text.replace(
                    "The Killing equation ensures that g_{7μ} = 0",
                    "The Killing equation ensures that ∂₇g_{μν} = 0"
                )
                fixes_applied += 1
                print(f"  FIX 1a' applied at paragraph {i}: Lemma statement variant")

    if "Killing equation" in text and "forces g_{" in text and i > 2900:
        for run in p.runs:
            rt = run.text
            if "forces g_{7μ} = 0" in rt or "forces g_7" in rt:
                old = rt
                rt = rt.replace("forces g_{7μ} = 0 for μ ≠ 7",
                               "forces ∂₇g_{μν} = 0 for all μ, ν — i.e., the metric is independent of the harm coordinate x⁷")
                rt = rt.replace("implying orthogonality of Dimension 7 to all others",
                               "implying that the metric does not depend on the harm coordinate. Note: this gives metric independence, not orthogonality. Orthogonality (g_{7μ} = 0 for μ ≠ 7) would require the additional condition that ξ is hypersurface-orthogonal (ξ ∧ dξ = 0). We adopt this as an additional modeling assumption: the harm dimension does not metrically mix with other moral dimensions under re-description. This is justified by the empirical finding that harm magnitude scales independently of categorical moral structure (§17.6)")
                if rt != old:
                    run.text = rt
                    fixes_applied += 1
                    print(f"  FIX 1b applied at paragraph {i}: Killing proof")

    # ═══════════════════════════════════════════════════════════════
    # FIX 2: Convex cone claim (paragraph ~3028)
    # ═══════════════════════════════════════════════════════════════
    if "Adm(M) is a convex cone" in text:
        for run in p.runs:
            if "convex cone" in run.text:
                run.text = run.text.replace(
                    "Adm(M) is a convex cone: if g, g' ∈ Adm(M) and α, β > 0, then αg + βg' ∈ Adm(M)",
                    "Adm(M) is path-connected and star-shaped with respect to the flat metric δ: for any g ∈ Adm(M) and t ∈ [0,1], the convex combination (1-t)δ + tg ∈ Adm(M). [Note: Adm(M) is not a convex cone in general, because the non-degeneracy and boundary constraints are not preserved under arbitrary positive linear combinations. The weaker star-shaped property suffices for the Structured Pluralism theorem below, which requires only that admissible perturbations can be constructed along line segments from δ.]"
                )
                fixes_applied += 1
                print(f"  FIX 2 applied at paragraph {i}: convex cone -> star-shaped")

    # ═══════════════════════════════════════════════════════════════
    # FIX 3: CHSH / SU(2) exclusion (paragraphs ~3519-3520)
    # ═══════════════════════════════════════════════════════════════
    if "Corollary 12.1" in text and "Exclusion of Non-Abelian" in text:
        for run in p.runs:
            if "inconsistent with any non-abelian continuous gauge component" in run.text:
                run.text = run.text.replace(
                    "inconsistent with any non-abelian continuous gauge component. In particular, SU(2) is excluded.",
                    "consistent with an abelian continuous gauge component and inconsistent with the maximal-violation predictions of non-abelian models at the tested sample size. In particular, the data do not exhibit the Tsirelson-bound violations that SU(2) with maximally entangled states would predict."
                )
                fixes_applied += 1
                print(f"  FIX 3a applied at paragraph {i}: CHSH corollary statement")

    if "SU(2) predicts Tsirelson-bound violations" in text and "By contrapositive" in text:
        for run in p.runs:
            if "By contrapositive" in run.text:
                run.text = run.text.replace(
                    "By contrapositive, any non-abelian continuous component (which would produce violations above the classical bound) is excluded by the data.",
                    "This does not constitute a proof that no non-abelian continuous component exists — SU(2) can produce |S| ≤ 2 for non-maximally-entangled or separable states. The result excludes the specific prediction of maximal-violation behavior under SU(2), not SU(2) itself. The data are consistent with the abelian model D₄ × U(1)_H and inconsistent with models that predict |S| > 2 at the tested configurations. A stronger exclusion would require testing across a larger space of measurement settings and establishing that no SU(2) representation produces the observed correlation structure."
                )
                fixes_applied += 1
                print(f"  FIX 3b applied at paragraph {i}: CHSH proof")

    # Also fix the summary line
    if "CHSH excludes non-abelian alternatives" in text:
        for run in p.runs:
            if "CHSH excludes non-abelian alternatives" in run.text:
                run.text = run.text.replace(
                    "CHSH excludes non-abelian alternatives.",
                    "CHSH data consistent with abelian model; do not exhibit non-abelian maximal violations."
                )
                fixes_applied += 1
                print(f"  FIX 3c applied at paragraph {i}: summary line")

    # ═══════════════════════════════════════════════════════════════
    # FIX 4: Conservation of harm — add modeling caveat (para ~3460)
    # ═══════════════════════════════════════════════════════════════
    if "This is not a metaphor. It is a formal result" in text and "Noether employed" in text:
        for run in p.runs:
            if "using the same mathematical structure that Noether employed" in run.text:
                run.text = run.text.replace(
                    "using the same mathematical structure that Noether employed. The parallel between physics and ethics is structural: both domains admit an action principle, both have symmetries, and both yield conservation laws from those symmetries.",
                    "using the same mathematical structure that Noether employed. The parallel between physics and ethics is structural: both domains admit an action principle, both have symmetries, and both yield conservation laws from those symmetries.\n\nAn important caveat: the conservation of harm is a modeling theorem, not an empirical discovery in the same sense as conservation of energy. In physics, the symmetry (time-translation invariance) is independently motivated by the observed laws of nature, and the conservation law (energy) is a surprising consequence. Here, the conserved quantity (harm) is identified with the Noether charge associated with a symmetry (BIP) that is itself a modeling assumption. The conservation law follows deductively from the framework's axioms — it is not extracted from data in the way Noether's theorem extracts conservation of energy from temporal symmetry. The empirical content lies in the axioms: does moral evaluation admit a Lagrangian structure? Is BIP a valid symmetry? If both hold, conservation follows as a structural consequence. The value of the result is its testability (the cross-lingual invariance prediction of §12.8) and its operational utility (the harm ledger of §12.6), not its surprisingness."
                )
                fixes_applied += 1
                print(f"  FIX 4 applied at paragraph {i}: conservation caveat")

    # ═══════════════════════════════════════════════════════════════
    # FIX 5: A* intractability — needs to be found in Chapter 11
    # Look for the bounded-suboptimality / intractability claims
    # ═══════════════════════════════════════════════════════════════
    if "exact moral geodesic planning is intractable" in text.lower() or \
       ("intractable" in text.lower() and "moral" in text.lower() and "optimization" in text.lower() and i > 2600):
        for run in p.runs:
            if "intractable" in run.text.lower():
                old = run.text
                # Add formal complexity reduction
                run.text = run.text + "\n\n[Formal complexity result.] Define MORAL-PATHFIND as the decision problem: given a weighted simplicial complex C with n vertices, d-dimensional attribute vectors with Mahalanobis edge weights, and k ≥ 1 attribute-boundary constraints (edges that cross a constraint boundary incur penalty βₖ), does there exist a path from s to G with total cost ≤ B? MORAL-PATHFIND is NP-hard by reduction from the Weight-Constrained Shortest Path problem (Garey & Johnson, 1979; Handler & Zang, 1980), which is NP-hard for k ≥ 1 side constraints. The reduction maps each side constraint to a boundary penalty β_k and each edge weight to the corresponding Mahalanobis distance. Since MORAL-PATHFIND is NP-hard, exact optimization on the moral manifold with boundary constraints is computationally intractable in the worst case, and the bounded-suboptimality guarantees of Weighted A* (Pohl, 1970) apply by direct instantiation — not by analogy — to the moral search problem."
                if run.text != old:
                    fixes_applied += 1
                    print(f"  FIX 5 applied at paragraph {i}: A* formalization")


# ═══════════════════════════════════════════════════════════════
# Update version number in title
# ═══════════════════════════════════════════════════════════════
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if "v1.20" in run.text:
            run.text = run.text.replace("v1.20", "v1.21")
            print(f"  Version updated at paragraph {i}")

print(f"\nTotal fixes applied: {fixes_applied}")
print(f"Saving: {DST}")
doc.save(DST)
print("Done!")
